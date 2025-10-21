# -*- coding: utf-8 -*-
"""
Gera dossiês de testemunhas (DOCX) a partir de um Excel em português.
Mudanças solicitadas:
- Removido: AdvogadoOuEscritorio
- Removido: Comportamento
- Renomeado: DocumentosOuExibicoes -> Provas

Requer: pandas, python-docx
"""

import os
import re
import sys
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.shared import OxmlElement, qn

# ------------ Configuração ------------
INPUT_XLSX_DEFAULT = "testemunhas.xlsx"  # caminho do Excel (padrão)
SAIDA_DIR_DEFAULT = "dossies"            # pasta de saída para os .docx (padrão)
FONTE = "Calibri"                        # ou "Times New Roman"
TAMANHO_FONTE = 11
TAMANHO_TITULO = 14
MARGENS_POL = 0.75                       # 0.75 polegadas ~ 19 mm
LINHAS_ANOTACOES = 15                    # número de linhas no bloco de anotações
ALTURA_LINHA_ANOTACOES = 0.4             # altura de cada linha em polegadas
# -------------------------------------

# Nomes de colunas esperados no Excel (português) após as mudanças
COLS = {
    "NomeTestemunha": "NomeTestemunha",
    "FuncaoOuRelacao": "FuncaoOuRelacao",
    "IntimadaPor": "IntimadaPor",
    "DataAudiencia": "DataAudiencia",
    "Processo": "Processo",
    "Contato": "Contato",
    "TestemunhosAnteriores": "TestemunhosAnteriores",
    "FatosChave": "FatosChave",
    "Estrategia": "Estrategia",
    "PerguntasPreparadas": "PerguntasPreparadas",
    "PontosCriticos": "PontosCriticos",
    "Provas": "Provas",
    "Etiquetas": "Etiquetas",
}

# Campos com múltiplas linhas
CAMPOS_LISTA = [
    COLS["FatosChave"],
    COLS["Estrategia"],
    COLS["PerguntasPreparadas"],
    COLS["PontosCriticos"],
    COLS["Provas"],
    COLS["TestemunhosAnteriores"],
]

# Cabeçalho (duas colunas)
CABECALHO_ESQ = [
    ("TESTEMUNHA", COLS["NomeTestemunha"]),
    ("Função/Relacao", COLS["FuncaoOuRelacao"]),
    ("Intimada por", COLS["IntimadaPor"]),
]
CABECALHO_DIR = [
    ("Processo", COLS["Processo"]),
    ("Data da audiência", COLS["DataAudiencia"]),
    # Removido "Advogado/Escritório"
]

def dividir_linhas(valor):
    """Quebra por nova linha ou ponto e vírgula; limpa espaços vazios."""
    if pd.isna(valor):
        return []
    txt = str(valor).strip()
    if not txt:
        return []
    partes = re.split(r'(?:\r?\n|;)', txt)
    return [p.strip() for p in partes if p.strip()]

def ajustar_margens(section, polegadas):
    section.top_margin = Inches(polegadas)
    section.bottom_margin = Inches(polegadas)
    section.left_margin = Inches(polegadas)
    section.right_margin = Inches(polegadas)

def estilo_default(doc):
    s = doc.styles['Normal']
    s.font.name = FONTE
    s.font.size = Pt(TAMANHO_FONTE)

def adiciona_titulo(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto.upper())
    run.bold = True
    run.font.size = Pt(TAMANHO_TITULO)
    return p

def adiciona_bullets(doc, itens):
    for item in itens:
        p = doc.add_paragraph(item, style=None)
        p.style = doc.styles['List Bullet']
        p.paragraph_format.space_after = Pt(2)

def adiciona_perguntas_como_lista(doc, perguntas_texto):
    """Formata perguntas como itens de lista numerada ou com bullets."""
    if not perguntas_texto or pd.isna(perguntas_texto):
        return
    
    # Divide o texto em linhas individuais
    linhas = dividir_linhas(perguntas_texto)
    if not linhas:
        return
    
    # Adiciona cada pergunta como um item de lista
    for pergunta in linhas:
        if pergunta.strip():  # Só adiciona se não for vazia
            p = doc.add_paragraph(pergunta.strip(), style=None)
            p.style = doc.styles['List Bullet']
            p.paragraph_format.space_after = Pt(2)

def bordas_tabela(tbl, cor="C0C0C0"):
    """Aplica bordas leves na tabela."""
    _tbl = tbl._tbl
    tblPr = _tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        _tbl.append(tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), cor)
        tblBorders.append(tag)
    tblPr.append(tblBorders)

def adiciona_cabecalho(doc, linha):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    bordas_tabela(tbl, cor="C0C0C0")

    cel_esq = tbl.cell(0, 0)
    cel_dir = tbl.cell(0, 1)

    # Nome grande no topo-esquerda
    nome = pegar(linha, COLS["NomeTestemunha"])
    p_nome = cel_esq.paragraphs[0]
    p_nome.clear()
    r = p_nome.add_run(nome)
    r.bold = True
    r.font.size = Pt(TAMANHO_TITULO + 2)

    # Bloco esquerdo
    for rotulo, chave in CABECALHO_ESQ[1:]:  # já usamos o nome como 1ª linha
        valor = pegar(linha, chave)
        pr = cel_esq.add_paragraph()
        run = pr.add_run(f"{rotulo}: ")
        run.bold = True
        pr.add_run(valor)

    # Bloco direito
    for rotulo, chave in CABECALHO_DIR:
        valor = pegar(linha, chave)
        pr = cel_dir.add_paragraph()
        run = pr.add_run(f"{rotulo}: ")
        run.bold = True
        pr.add_run(valor)

def adiciona_bloco_anotacoes(doc, linhas=12, altura_linha=0.4):
    adiciona_titulo(doc, "Anotações em audiência")
    tbl = doc.add_table(rows=linhas, cols=1)
    tbl.style = 'Table Grid'
    for i in range(linhas):
        row = tbl.rows[i]
        row.height = Inches(altura_linha)  # Define altura da linha
        cell = tbl.cell(i, 0)
        p = cell.paragraphs[0]
        p.add_run(" ")
    bordas_tabela(tbl, cor="DDDDDD")

def formatar_data(valor):
    if pd.isna(valor) or str(valor).strip() == "":
        return ""
    if isinstance(valor, (pd.Timestamp, datetime)):
        return valor.strftime("%Y-%m-%d")
    try:
        return pd.to_datetime(valor).strftime("%Y-%m-%d")
    except Exception:
        return str(valor)

def pegar(linha_dict, chave, default=""):
    v = linha_dict.get(chave, default)
    if chave == COLS["DataAudiencia"]:
        return formatar_data(v)
    # Trata NaN/strings vazias com cuidado
    try:
        if pd.isna(v):
            return ""
    except Exception:
        pass
    return str(v).strip()

def construir_doc(linha, caminho_saida):
    doc = Document()
    estilo_default(doc)
    ajustar_margens(doc.sections[0], MARGENS_POL)

    # Cabeçalho
    adiciona_cabecalho(doc, linha)

    # Seções (ordem e títulos em PT-BR)
    secoes = [
        ("Fatos-chave", COLS["FatosChave"]),
        ("Estratégia", COLS["Estrategia"]),
        ("Pontos críticos / Impeachment", COLS["PontosCriticos"]),
        ("Provas", COLS["Provas"]),
        ("Testemunhos anteriores", COLS["TestemunhosAnteriores"]),
    ]
    for titulo, chave in secoes:
        itens = dividir_linhas(linha.get(chave, ""))
        if itens:
            adiciona_titulo(doc, titulo)
            adiciona_bullets(doc, itens)
    
    # Seção especial para perguntas preparadas (como lista)
    perguntas = linha.get(COLS["PerguntasPreparadas"], "")
    if perguntas and not pd.isna(perguntas) and str(perguntas).strip():
        adiciona_titulo(doc, "Perguntas preparadas")
        adiciona_perguntas_como_lista(doc, perguntas)

    # Bloco de anotações
    adiciona_bloco_anotacoes(doc, linhas=LINHAS_ANOTACOES, altura_linha=ALTURA_LINHA_ANOTACOES)

    # Rodapé: etiquetas
    etiquetas = pegar(linha, COLS["Etiquetas"])
    if etiquetas:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"Etiquetas: {etiquetas}")
        run.italic = True

    doc.save(caminho_saida)

def main():
    # Aceita argumentos da linha de comando
    input_xlsx = sys.argv[1] if len(sys.argv) > 1 else INPUT_XLSX_DEFAULT
    saida_dir = sys.argv[2] if len(sys.argv) > 2 else SAIDA_DIR_DEFAULT
    
    # Lembrete sobre delimitadores
    print("\n⚠️  LEMBRETE: Use ponto e vírgula (;) para separar itens em listas no Excel.")
    print("   Exemplo: 'Pergunta 1; Pergunta 2; Pergunta 3'\n")
    
    if not os.path.exists(input_xlsx):
        print(f"Erro: Arquivo não encontrado: {input_xlsx}")
        print(f"Uso: python {sys.argv[0]} [arquivo_excel] [pasta_saida]")
        sys.exit(1)
    
    os.makedirs(saida_dir, exist_ok=True)
    df = pd.read_excel(input_xlsx)
    df.columns = [c.strip() for c in df.columns]

    for i, r in df.iterrows():
        linha = r.to_dict()
        nome = pegar(linha, COLS["NomeTestemunha"]) or f"Testemunha_{i+1}"
        nome_seguro = re.sub(r'[^A-Za-z0-9._ -À-ÿ]+', '_', nome).strip()
        saida = os.path.join(saida_dir, f"{nome_seguro}.docx")
        construir_doc(linha, saida)
        print(f"✓ {saida}")

if __name__ == "__main__":
    main()
